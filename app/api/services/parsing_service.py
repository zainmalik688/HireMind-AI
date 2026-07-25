import io
import re
from typing import Any

import fitz  # PyMuPDF
from docx import Document
import pytesseract
from PIL import Image

from app.api.schemas import (
    AwardItem,
    CertificationItem,
    FileValidationResult,
    InterestItem,
    LanguageItem,
    ParsedResumeData,
    PublicationItem,
    ReferenceItem,
    ResumeQualityCheck,
)


class ResumeParsingEngine:

    # Line-start Header Patterns (All 6 core secondary sections)
    CERTIFICATION_HEADER_PATTERN = r'(?i)^\s*(?:3\.\s*|•\s*)?(?:certifications?|licenses?(?:\s*&\s*certifications?)?|credentials?|professional\s+certifications?|courses?\s*&\s*certifications?)\b'
    LANGUAGE_HEADER_PATTERN = r'(?i)^\s*(?:3\.\s*|•\s*)?(?:languages?|languages?\s*spoken|spoken\s*languages?|language\s*proficiency)\b'
    AWARD_HEADER_PATTERN = r'(?i)^\s*(?:3\.\s*|•\s*)?(?:awards?|honors?|awards?\s*&\s*honors?|honors?\s*&\s*awards?|achievements?|academic\s+honors?)\b'
    PUBLICATION_HEADER_PATTERN = r'(?i)^\s*(?:3\.\s*|•\s*)?(?:publications?|research\s+papers?|papers?|articles?|journals?|conference\s+proceedings?)\b'
    INTEREST_HEADER_PATTERN = r'(?i)^\s*(?:3\.\s*|•\s*)?(?:interests?|hobbies?|personal\s+interests?|extracurricular\s+activities?|hobbies\s*&\s*interests?)\b'
    REFERENCE_HEADER_PATTERN = r'(?i)^\s*(?:3\.\s*|•\s*)?(?:references?|professional\s+references?|referees?)\b'

    # Universal Header Detector for Section Boundary Splitting
    ALL_HEADERS_PATTERN = r'(?im)^\s*(?:[0-9]+\.|\u2022|-|\*)*\s*(?:certifications?|licenses?|credentials?|languages?|awards?|honors?|achievements?|publications?|research\s+papers?|interests?|hobbies?|references?|referees?|experience|education|skills|projects|summary|work\s+experience)\b.*$'

    @staticmethod
    def _clean_extracted_text(text: str) -> str:
        """
        Normalizes extracted resume text across PDF and DOCX formats.
        Fixes split-word hyphens, standardizes line breaks, and preserves structural sections.
        """
        if not text:
            return ""

        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r'[•▪‣⁃●\u2022\u2023\u2043\u2219]', '-', text)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)

        return text.strip()

    @classmethod
    def _extract_section_blocks(cls, text: str) -> dict[str, str]:
        """
        Splits clean resume text into section blocks based on header pattern matching.
        """
        section_map: dict[str, str] = {}
        lines = text.split('\n')
        current_section = None
        buffer: list[str] = []

        patterns = {
            "certifications": cls.CERTIFICATION_HEADER_PATTERN,
            "languages": cls.LANGUAGE_HEADER_PATTERN,
            "awards": cls.AWARD_HEADER_PATTERN,
            "publications": cls.PUBLICATION_HEADER_PATTERN,
            "interests": cls.INTEREST_HEADER_PATTERN,
            "references": cls.REFERENCE_HEADER_PATTERN,
        }

        for line in lines:
            matched_key = None
            for key, pattern in patterns.items():
                if re.search(pattern, line):
                    matched_key = key
                    break

            if matched_key:
                if current_section and buffer:
                    section_map[current_section] = "\n".join(buffer).strip()
                    buffer = []
                current_section = matched_key
            elif current_section:
                # Stop capturing if another general header is encountered
                if re.match(cls.ALL_HEADERS_PATTERN, line) and not any(re.search(p, line) for p in patterns.values()):
                    section_map[current_section] = "\n".join(buffer).strip()
                    current_section = None
                    buffer = []
                else:
                    buffer.append(line)

        if current_section and buffer:
            section_map[current_section] = "\n".join(buffer).strip()

        return section_map

    @classmethod
    def _parse_structured_sections(cls, text: str) -> dict[str, list[Any]]:
        """
        Transforms raw extracted text section blocks into populated Pydantic models.
        """
        blocks = cls._extract_section_blocks(text)
        
        certifications = [
            CertificationItem(title=line.strip("- ").strip())
            for line in blocks.get("certifications", "").split("\n")
            if line.strip("- ").strip()
        ]

        languages = [
            LanguageItem(
                language=line.split("-")[0].strip("- ").strip(),
                proficiency=line.split("-")[1].strip() if "-" in line else None
            )
            for line in blocks.get("languages", "").split("\n")
            if line.strip("- ").strip()
        ]

        awards = [
            AwardItem(title=line.strip("- ").strip())
            for line in blocks.get("awards", "").split("\n")
            if line.strip("- ").strip()
        ]

        publications = [
            PublicationItem(title=line.strip("- ").strip())
            for line in blocks.get("publications", "").split("\n")
            if line.strip("- ").strip()
        ]

        interests = [
            InterestItem(name=line.strip("- ").strip())
            for line in blocks.get("interests", "").split("\n")
            if line.strip("- ").strip()
        ]

        references = [
            ReferenceItem(name=line.strip("- ").strip())
            for line in blocks.get("references", "").split("\n")
            if line.strip("- ").strip()
        ]

        return {
            "certifications": certifications,
            "languages": languages,
            "awards": awards,
            "publications": publications,
            "interests": interests,
            "references": references,
        }

    @staticmethod
    def parse_pdf(content: bytes) -> tuple[str, bool, dict[str, Any]]:
        doc = fitz.open(stream=content, filetype="pdf")
        extracted_text = ""
        extracted_links: list[str] = []
        is_scanned = False
        metadata: dict[str, Any] = dict(doc.metadata or {})

        for page in doc:
            extracted_text += f"{page.get_text('text')}\n\n"
            for link in page.get_links():
                uri = link.get("uri")
                if uri and uri.startswith(("http://", "https://", "mailto:")) and uri not in extracted_links:
                    extracted_links.append(uri)

        # Fallback to OCR if page has minimal/no raw text (Scanned PDF)
        if len(extracted_text.strip()) < 50:
            is_scanned = True
            extracted_text = ""
            for page in doc:
                pix = page.get_pixmap()
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                extracted_text += f"{pytesseract.image_to_string(img)}\n\n"

        if extracted_links:
            extracted_text += "\n\n--- EXTRACTED HYPERLINKS ---\n" + "\n".join(extracted_links)

        return extracted_text, is_scanned, metadata

    @staticmethod
    def parse_docx(content: bytes) -> tuple[str, dict[str, Any]]:
        doc = Document(io.BytesIO(content))
        full_text: list[str] = []
        extracted_links: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(f"\n{text}\n" if para.style.name.startswith('Heading') else text)

        for table in doc.tables:
            for row in table.rows:
                if row_text := [cell.text.strip() for cell in row.cells if cell.text.strip()]:
                    full_text.append(" | ".join(row_text))

        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype and rel.target_ref:
                url = rel.target_ref
                if url.startswith(("http://", "https://", "mailto:")) and url not in extracted_links:
                    extracted_links.append(url)

        combined_text = "\n".join(full_text)
        if extracted_links:
            combined_text += "\n\n--- EXTRACTED HYPERLINKS ---\n" + "\n".join(extracted_links)

        metadata: dict[str, Any] = {
            "author": doc.core_properties.author or "",
            "title": doc.core_properties.title or ""
        }
        return combined_text, metadata

    @staticmethod
    def parse_txt(content: bytes) -> tuple[str, dict[str, Any]]:
        return content.decode("utf-8", errors="ignore"), {}

    @classmethod
    def process_document(cls, file_name: str, validation_info: FileValidationResult, content: bytes) -> ParsedResumeData:
        raw_text, is_scanned, metadata = "", False, {}

        match validation_info.file_type.lower():
            case "pdf":
                raw_text, is_scanned, metadata = cls.parse_pdf(content)
                validation_info.is_scanned = is_scanned
            case "docx":
                raw_text, metadata = cls.parse_docx(content)
            case "txt":
                raw_text, metadata = cls.parse_txt(content)

        cleaned_text = cls._clean_extracted_text(raw_text)
        words = cleaned_text.split()
        word_count = len(words)

        resume_keywords = {
            "experience", "education", "skills", "projects", "summary", "work", 
            "university", "certifications", "courses", "languages", "awards", 
            "honors", "publications", "research", "interests", "hobbies", 
            "references", "referees"
        }
        matches = sum(1 for word in words if word.lower() in resume_keywords)
        confidence = min(round((matches / 4.0), 2), 1.0) if word_count > 30 else 0.1

        quality_info = ResumeQualityCheck(
            is_resume=(confidence >= 0.4 and word_count >= 50),
            confidence_score=confidence,
            word_count=word_count,
            char_count=len(cleaned_text),
            quality_notes=[
                "Sufficient word length detected." if word_count >= 50 else "Document is extremely short.",
                "High density of resume section keywords." if confidence >= 0.6 else "Low density of standard resume headings."
            ]
        )

        # Parse structured sections into Pydantic schema lists
        parsed_sections = cls._parse_structured_sections(cleaned_text)

        return ParsedResumeData(
            raw_text=raw_text,
            cleaned_text=cleaned_text,
            file_name=file_name,
            file_type=validation_info.file_type,
            metadata=metadata,
            validation_info=validation_info,
            quality_info=quality_info,
            certifications=parsed_sections["certifications"],
            languages=parsed_sections["languages"],
            awards=parsed_sections["awards"],
            publications=parsed_sections["publications"],
            interests=parsed_sections["interests"],
            references=parsed_sections["references"]
        )