import io
import logging
import re
from pathlib import Path
from typing import Any
import fitz  # PyMuPDF
import docx
from PIL import Image
import pytesseract

from app.api.services.ats_parsing_checker import (
    check_pdf_parsing_issues,
    check_docx_parsing_issues,
)

logger = logging.getLogger(__name__)


def is_corrupted_text_content(text: str) -> bool:
    """
    Flags decoded text that doesn't actually look like readable text --
    e.g. a binary file (image, executable, etc.) renamed to .txt, which
    decodes "successfully" under latin-1 but isn't real content. Uses the
    same signal `file`/`git` use to tell binary from text: NUL bytes, or
    a high ratio of non-printable control characters.
    """
    if not text:
        return False
    if "\x00" in text:
        return True
    control_chars = sum(
        1 for ch in text
        if (ord(ch) < 32 and ch not in "\n\r\t") or ord(ch) == 127
    )
    return (control_chars / len(text)) > 0.05


def extract_text_from_txt(file_input: bytes | Path | str) -> dict[str, Any]:
    """
    Extract raw text from a plain text file (.txt).
    Accepts raw bytes, Path, or file path string.
    """
    try:
        if isinstance(file_input, bytes):
            # Try UTF-8 decoding with fallback to latin-1
            try:
                extracted_text = file_input.decode("utf-8")
            except UnicodeDecodeError:
                extracted_text = file_input.decode("latin-1")
        else:
            path = Path(file_input)
            extracted_text = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as err:
        raise ValueError(f"Failed to read TXT document: {str(err)}") from err

    if is_corrupted_text_content(extracted_text):
        raise ValueError("The TXT document appears to be corrupted or unreadable.")

    extracted_text = extracted_text.strip()
    word_count = len(extracted_text.split())

    if not extracted_text:
        raise ValueError("The TXT document contains no readable text.")

    # Extract embedded hyperlinks/URLs via regex
    url_pattern = r'https?://[^\s<>"]+|www\.[^\s<>"]+'
    extracted_links = list(set(re.findall(url_pattern, extracted_text)))

    final_text = extracted_text
    if extracted_links:
        links_block = "\n\n--- EXTRACTED HYPERLINKS ---\n" + "\n".join(extracted_links)
        final_text += links_block

    return {
        "text": final_text,
        "raw_text": extracted_text,
        "is_scanned": False,
        "word_count": word_count,
        "image_count": 0,
        "page_count": 1,
        "extracted_links": extracted_links,
    }


def extract_text_from_pdf(file_input: bytes | Path | str) -> dict[str, Any]:
    """
    Extract raw text, embedded URI annotations, metadata, and scanned state from a PDF.
    Accepts raw bytes, Path, or file path string.
    """
    text_content: list[str] = []
    extracted_links: list[str] = []
    total_images = 0

    try:
        if isinstance(file_input, bytes):
            doc = fitz.open(stream=file_input, filetype="pdf")
        else:
            doc = fitz.open(file_input)

        with doc:
            page_count = len(doc)
            for page in doc:
                # 1. Extract body text
                text = page.get_text("text")
                if text:
                    text_content.append(text)

                # 2. Count images per page to evaluate scanned likelihood
                total_images += len(page.get_images(full=True))

                # 3. Extract embedded hyperlink annotations
                links = page.get_links()
                for link in links:
                    uri = link.get("uri")
                    if uri and uri.startswith(("http://", "https://", "mailto:")):
                       if uri not in extracted_links:
                            extracted_links.append(uri)

            # Deterministic structural ATS evidence, computed on this same
            # already-open document -- no second fitz.open(). Isolated in
            # its own boundary so a checker failure never blocks extraction.
            try:
                parsing_issues = check_pdf_parsing_issues(doc)
            except Exception as check_err:
                logger.info(f"[ATS Parsing Check] PDF check failed: {type(check_err).__name__}")
                parsing_issues = []

            extracted_text = "\n".join(text_content).strip()
            word_count = len(extracted_text.split())

            # Detect scanned document: low readable words (< 30) while containing images OR zero text entirely
            is_scanned = (word_count < 30) and (total_images > 0 or word_count == 0)

            if is_scanned:
                try:
                    print("Document OCR Fallback Executed")
                    extracted_text = ""
                    for page in doc:
                        pix = page.get_pixmap()
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        extracted_text += f"{pytesseract.image_to_string(img)}\n\n"
                    word_count = len(extracted_text.split())
                    is_scanned = False  # Reset is_scanned after OCR
                except Exception as ocr_err:
                    raise ValueError(f"OCR failed: {str(ocr_err)}") from ocr_err

    except ValueError:
        raise
    except Exception as err:
        raise ValueError(f"Failed to parse PDF document: {str(err)}") from err

    # 4. Append discovered URIs to the text stream so downstream entity extractors pick them up
    final_text = extracted_text
    if extracted_links:
        links_block = "\n\n--- EXTRACTED HYPERLINKS ---\n" + "\n".join(extracted_links)
        final_text += links_block

    return {
        "text": final_text,
        "raw_text": extracted_text,
        "is_scanned": is_scanned,
        "word_count": word_count,
        "image_count": total_images,
        "page_count": page_count,
        "extracted_links": extracted_links,
        "parsing_issues": parsing_issues,
    }


def extract_text_from_docx(file_input: bytes | Path | str) -> dict[str, Any]:
    """
    Extract raw text and embedded hyperlinks from a DOCX document.
    Accepts raw bytes, Path, or file path string.
    """
    extracted_links: list[str] = []
    paragraphs: list[str] = []

    try:
        if isinstance(file_input, bytes):
            doc = docx.Document(io.BytesIO(file_input))
        else:
            doc = docx.Document(file_input)

        # 1. Extract paragraphs text
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())

        # 2. Extract embedded relationship links from document XML
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype and rel.target_ref:
                url = rel.target_ref
                if url.startswith(("http://", "https://", "mailto:")) and url not in extracted_links:
                    extracted_links.append(url)

    except Exception as err:
        raise ValueError(f"Failed to parse DOCX document: {str(err)}") from err

    # Deterministic structural ATS evidence, computed on this same
    # already-open document -- no second docx.Document(...). Isolated in
    # its own boundary so a checker failure never blocks extraction.
    try:
        parsing_issues = check_docx_parsing_issues(doc)
    except Exception as check_err:
        logger.info(f"[ATS Parsing Check] DOCX check failed: {type(check_err).__name__}")
        parsing_issues = []

    extracted_text = "\n".join(paragraphs).strip()
    word_count = len(extracted_text.split())

    if not extracted_text:
        raise ValueError("The DOCX document contains no readable text.")

    final_text = extracted_text
    if extracted_links:
        links_block = "\n\n--- EXTRACTED HYPERLINKS ---\n" + "\n".join(extracted_links)
        final_text += links_block

    return {
        "text": final_text,
        "raw_text": extracted_text,
        "is_scanned": False,
        "word_count": word_count,
        "image_count": 0,
        "page_count": 1,
        "extracted_links": extracted_links,
        "parsing_issues": parsing_issues,
    }


def extract_text_from_file(file_input: bytes | Path | str, filename: str = "") -> dict[str, Any]:
    """
    Extract text, metadata, and scanned metrics from supported document formats (.pdf, .docx, .txt).
    Accepts raw bytes with filename OR a Path / string object.
    """
    if isinstance(file_input, (str, Path)):
        path = Path(file_input)
        suffix = path.suffix.lower()
    else:
        suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_input)
    elif suffix == ".docx":
        return extract_text_from_docx(file_input)
    elif suffix == ".txt":
        return extract_text_from_txt(file_input)
    else:
        raise ValueError(f"Unsupported file format: '{suffix}'. Supported formats: .pdf, .docx, .txt")
