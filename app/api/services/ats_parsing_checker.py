import re
from docx.oxml.ns import qn
from app.api.schemas import ATSParsingIssue

# --- Config: isolated here so thresholds can be tuned without touching logic ---
MIN_BLOCKS_PER_COLUMN = 3
MIN_COLUMN_HEIGHT_RATIO = 0.24
X_CLUSTER_TOLERANCE = 25
MIN_DIGITS_FOR_PHONE = 7

BULLET_PATTERN = re.compile(r'[\uf0b7\uf06f\ue000-\uf8ff]')
EMAIL_PATTERN = re.compile(r'[\w.+-]+@[\w-]+\.\w+')
PHONE_CANDIDATE_PATTERN = re.compile(r'\+?[\d\s\-()]{8,}')

# Wingdings-style bullet glyphs and WinAnsi middle dot (U+00B7) share the
# same underlying byte value, so PDF extraction commonly normalizes a
# dingbat bullet to "·" instead of preserving a private-use codepoint.
# Only the line-leading position is treated as a bullet marker; "·" used
# inline (e.g. "City · email · phone") is ordinary punctuation and is
# never matched by this pattern.
MIDDOT_LINE_START_PATTERN = re.compile(r'^[ \t]*·', re.MULTILINE)
MIN_MIDDOT_BULLET_LINES = 2

# Approved, conservative descriptions (Step 2 - Change 3). These state the
# structural fact the detector can actually prove, not universal ATS behavior.
DESCRIPTIONS = {
    "TABLE": "Table detected — table-based layouts may reduce ATS parsing reliability.",
    "HEADER_FOOTER_TEXT": "Contact information detected in header/footer — this placement may not be reliably parsed by ATS.",
    "TEXT_BOX": "Text box detected — content inside text boxes may not be reliably parsed by ATS.",
    "MULTI_COLUMN": "Layout suggests multiple columns — this may affect ATS reading order.",
    "NONSTANDARD_BULLETS": "Icon or symbol bullets detected — these may render as garbled characters in some ATS.",
}

# Known icon/dingbat font families whose characters only make visual sense
# inside that font (arrows, checkmarks, stars, etc). Word's own default
# bullet uses the Symbol font and is intentionally NOT in this list.
NONSTANDARD_BULLET_FONTS = {"Wingdings", "Wingdings 2", "Wingdings 3", "Webdings", "Marlett"}


def _has_contact_info(text: str) -> bool:
    if EMAIL_PATTERN.search(text):
        return True
    for candidate in PHONE_CANDIDATE_PATTERN.findall(text):
        if sum(c.isdigit() for c in candidate) >= MIN_DIGITS_FOR_PHONE:
            return True
    return False


def _bullet_font(paragraph, numbering) -> tuple[str | None, str | None]:
    """Resolves a paragraph's OOXML bullet level to (numFmt, font).
    Falls back to the paragraph's style numPr since Word's built-in
    List Bullet/List Number styles carry numPr on the style, not the paragraph.
    """
    numPr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
    if numPr is None:
        style_pPr = paragraph.style.element.find(qn('w:pPr'))
        numPr = style_pPr.find(qn('w:numPr')) if style_pPr is not None else None
    if numPr is None or numPr.numId is None:
        return None, None

    ilvl = numPr.ilvl.val if numPr.ilvl is not None else 0
    num = numbering.find(f'{qn("w:num")}[@{qn("w:numId")}="{numPr.numId.val}"]')
    abstract_id = num.find(qn('w:abstractNumId')).get(qn('w:val')) if num is not None else None
    lvl = (num.find(f'{qn("w:lvlOverride")}[@{qn("w:ilvl")}="{ilvl}"]/{qn("w:lvl")}')
           or numbering.find(f'{qn("w:abstractNum")}[@{qn("w:abstractNumId")}="{abstract_id}"]'
                              f'/{qn("w:lvl")}[@{qn("w:ilvl")}="{ilvl}"]')) if abstract_id else None
    if lvl is None:
        return None, None

    fmt_el = lvl.find(qn('w:numFmt'))
    rFonts = lvl.find(f'{qn("w:rPr")}/{qn("w:rFonts")}')
    font = (rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi'))) if rFonts is not None else None
    return (fmt_el.get(qn('w:val')) if fmt_el is not None else None), font


def check_docx_parsing_issues(doc) -> list[ATSParsingIssue]:
    issues = []

    if doc.tables:
        issues.append(ATSParsingIssue(
            issue_type="TABLE", severity="high", confidence="high",
            description=DESCRIPTIONS["TABLE"]
        ))

    hf_text = "\n".join(
        p.text for s in doc.sections
        for p in s.header.paragraphs + s.footer.paragraphs
    )
    if _has_contact_info(hf_text):
        issues.append(ATSParsingIssue(
            issue_type="HEADER_FOOTER_TEXT", severity="high", confidence="high",
            description=DESCRIPTIONS["HEADER_FOOTER_TEXT"]
        ))

    if '<w:txbxContent' in doc.element.xml:
        issues.append(ATSParsingIssue(
            issue_type="TEXT_BOX", severity="high", confidence="high",
            description=DESCRIPTIONS["TEXT_BOX"]
        ))

    # w:cols/w:num is the literal OOXML column count Word writes for
    # Format > Columns -- an explicit structural fact, not inferred layout.
    has_multi_column = False
    for s in doc.sections:
        cols = s._sectPr.find(qn('w:cols'))
        if cols is not None and int(cols.get(qn('w:num'), 1)) >= 2:
            has_multi_column = True
            break

    if has_multi_column:
        issues.append(ATSParsingIssue(
            issue_type="MULTI_COLUMN", severity="medium", confidence="high",
            description=DESCRIPTIONS["MULTI_COLUMN"]
        ))

    # A bullet is only "nonstandard" when its OOXML numbering level uses an
    # icon/dingbat font (Wingdings, Webdings, etc). Word's default Symbol-
    # font bullet, and any bullet with no font override, are left alone.
    numbering = doc.part.numbering_part.element
    has_nonstandard_bullets = False
    for paragraph in doc.paragraphs:
        fmt, font = _bullet_font(paragraph, numbering)
        if fmt == "bullet" and font in NONSTANDARD_BULLET_FONTS:
            has_nonstandard_bullets = True
            break

    if has_nonstandard_bullets:
        issues.append(ATSParsingIssue(
            issue_type="NONSTANDARD_BULLETS", severity="low", confidence="medium",
            description=DESCRIPTIONS["NONSTANDARD_BULLETS"]
        ))

    return issues


def _cluster_columns_sorted(blocks: list) -> list:
    """Order-independent clustering: sort by x0, merge points within tolerance."""
    if not blocks:
        return []
    sorted_blocks = sorted(blocks, key=lambda b: b[0])
    clusters = [{"x": sorted_blocks[0][0], "blocks": [sorted_blocks[0]]}]
    for b in sorted_blocks[1:]:
        if b[0] - clusters[-1]["x"] <= X_CLUSTER_TOLERANCE:
            clusters[-1]["blocks"].append(b)
        else:
            clusters.append({"x": b[0], "blocks": [b]})
    return clusters


def _count_real_columns(blocks: list, page_height: float) -> int:
    clusters = _cluster_columns_sorted(blocks)
    real_columns = 0
    for c in clusters:
        if len(c["blocks"]) < MIN_BLOCKS_PER_COLUMN:
            continue
        y_top = min(b[1] for b in c["blocks"])
        y_bottom = max(b[3] for b in c["blocks"])
        if page_height and (y_bottom - y_top) / page_height >= MIN_COLUMN_HEIGHT_RATIO:
            real_columns += 1
    return real_columns


def check_pdf_page_issues(fitz_page, page_number: int | None = None) -> list[ATSParsingIssue]:
    """Checks a single page. Caller aggregates across all pages.

    page_number is optional and 1-indexed. When supplied, it is attached to
    each generated issue's affected_pages. Existing callers that invoke this
    with just a page (no page_number) continue to work unchanged, and
    affected_pages is left as None in that case.
    """
    issues = []
    affected_pages = [page_number] if page_number is not None else None

    # PDF-native table evidence based on detected ruling/grid lines.
    # strategy="lines_strict" only considers actual drawn lines, so a
    # borderless fill or a lone decorative rule is unlikely to register --
    # this is detection evidence, not semantic proof a table exists (or,
    # on failure, proof one doesn't).
    #
    # find_tables() is a large third-party algorithm (ported from
    # pdfplumber) with no single documented exception type: it explicitly
    # raises ValueError/IndexError from its own geometry math, and can
    # also propagate whatever the underlying content-stream parsing
    # raises. Broad but narrowly scoped to this call only -- the same
    # pattern pdf_service.py already uses around the whole parsing-issue
    # check.
    has_table = False
    try:
        has_table = bool(fitz_page.find_tables(strategy="lines_strict").tables)
    except Exception:
        has_table = False  # no reliable TABLE evidence -- not proof of absence
    if has_table:
        issues.append(ATSParsingIssue(
            issue_type="TABLE", severity="high", confidence="high",
            affected_pages=affected_pages,
            description=DESCRIPTIONS["TABLE"]
        ))

    blocks = [b for b in fitz_page.get_text("blocks") if b[4].strip()]
    if _count_real_columns(blocks, fitz_page.rect.height) >= 2:
        issues.append(ATSParsingIssue(
            issue_type="MULTI_COLUMN", severity="medium", confidence="medium",
            affected_pages=affected_pages,
            description=DESCRIPTIONS["MULTI_COLUMN"]
        ))

    page_text = fitz_page.get_text()
    has_raw_pua_bullet = bool(BULLET_PATTERN.search(page_text))
    has_normalized_middot_bullets = len(MIDDOT_LINE_START_PATTERN.findall(page_text)) >= MIN_MIDDOT_BULLET_LINES
    if has_raw_pua_bullet or has_normalized_middot_bullets:
        issues.append(ATSParsingIssue(
            issue_type="NONSTANDARD_BULLETS", severity="low", confidence="medium",
            affected_pages=affected_pages,
            description=DESCRIPTIONS["NONSTANDARD_BULLETS"]
        ))

    return issues


def check_pdf_parsing_issues(fitz_doc) -> list[ATSParsingIssue]:
    """Loops every page, aggregates issue types, dedupes, and collects the
    accurate 1-indexed page numbers each issue was found on."""
    seen = {}
    for page_number, page in enumerate(fitz_doc, start=1):
        for issue in check_pdf_page_issues(page, page_number=page_number):
            if issue.issue_type not in seen:
                seen[issue.issue_type] = issue
            else:
                seen[issue.issue_type].affected_pages.append(page_number)

    results = []
    for issue in seen.values():
        if issue.affected_pages:
            issue.affected_pages = sorted(set(issue.affected_pages))
        results.append(issue)
    return results